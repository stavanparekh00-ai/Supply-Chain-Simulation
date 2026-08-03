"""
Generates the "Supply Chain Solver + Simulation — Project Design Brief" Word document.
Run with: python3 generate_design_brief.py
Output: Supply_Chain_Project_Design_Brief.docx (in the same folder)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---------- Style setup ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def add_title():
    title = doc.add_heading("Supply Chain Solver + Simulation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Project Design Brief — Discussion Summary")
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    run.bold = True
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(datetime.date.today().strftime("Compiled %B %d, %Y"))
    date_run.italic = True
    date_run.font.size = Pt(10)
    status_p = doc.add_paragraph()
    status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status_p.add_run(
        "Status: Design / discussion phase only — no code has been written yet."
    )
    status_run.bold = True
    status_run.font.size = Pt(11)
    doc.add_paragraph()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullet(text, bold_lead=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = para.add_run(bold_lead)
        r1.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def numbered(text):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)
    return para


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


# =========================================================================
add_title()

# ---------------------------------------------------------------------
h1("1. Project Overview & Structure")

p(
    "The project is deliberately split into two independent deliverables that share one "
    "common data schema, rather than being built as a single combined system."
)

h2("1.1 The Oracle Solver (Phase 1 — the portfolio piece)")
p(
    "A standalone, validated optimization engine. This is the resume/GitHub/interview "
    "deliverable, built and finished first, with no dependency on the simulation."
)

h2("1.2 The Simulation / Case Study (Phase 2 — the side project)")
p(
    "Friends/peers each play through an identical, structured scenario making period-by-period "
    "ordering decisions. Their decisions are then compared against the Oracle's mathematically "
    "optimal solution for that same scenario. Lower stakes, built after the solver is finished, "
    "iterated on at a relaxed pace."
)

h2("1.3 Why split rather than combine")
bullet("Matches the Oracle's own internal design principle (a clean, isolated function with a strict input/output contract) at the project level, not just in the code.")
bullet("De-risks the timeline — the solver alone is complete and shippable with no UI or game loop required.")
bullet("Produces two crisp, separate narratives instead of one diluted one: an optimization-rigor story and a decision-science story.")
bullet("The game becomes genuinely lower-pressure and can be iterated on based on real playthrough feedback, since it isn't the portfolio-critical piece.")

h2("1.4 The one shared dependency")
p(
    "The data schema (facility / supplier / disruption / period-state format) is designed once "
    "and shared by both projects, so the simulation can plug into the existing solver without "
    "rework later. This is the only coordination point between the two independent timelines."
)

# ---------------------------------------------------------------------
h1("2. Portfolio Positioning")

h2("2.1 Part 1 — The Solver (\"optimization depth\" credential)")
p("Framed as a decision-support / prescriptive-analytics engine, not a game backend.", italic=True)
bullet("Multi-supplier, multi-facility allocation under dedicated (non-pooled) capacity constraints.")
bullet("Tariff-aware procurement cost, modeling real cost exposure from international sourcing.")
bullet("Risk management enforced as a hard, tiered diversification constraint rather than a cost penalty — keeps the objective as pure cost-minimization while keeping risk mitigation explainable and auditable.")
bullet("Lead-time-aware checkpoint feasibility, not just an end-of-window target — prevents the solver from exploiting slow-but-cheap suppliers at the expense of near-term stockouts.")
bullet("Disruption-responsive capacity modeling — dedicated per facility-supplier capacity, cut unequally by priority rank during disruptions.")
bullet("Validated against hand-worked test cases, with any real bugs found during that process documented honestly (see Section 3.2 — do not pre-write this narrative before it actually happens).")

h2("2.2 Part 2 — The Human Case Study (\"decision science\" credential)")
p(
    "Framed as a genuine research question: how far do real human planners deviate from a "
    "mathematically optimal benchmark, when given identical information (no hindsight, no "
    "lookahead), and does that gap widen under supply disruption?",
    italic=True,
)
bullet("Identical-information benchmarking — the Oracle never sees anything the human participant didn't see at decision time. This is the single most important credibility sentence in the whole project.")
bullet("Decision-quality gap under uncertainty — compared period-by-period against what was optimal given the same forecast and visible state, not just final outcomes.")
bullet("Disruption sensitivity — a genuinely reportable finding once real data exists (e.g. \"the decision-quality gap widened by X% during disruption periods vs. stable periods\").")
bullet("Small-sample, explicitly-labeled qualitative study (roughly 5-10 playthroughs, friends/peers) — framed honestly as exploratory, not a rigorous statistical experiment. Honesty about scale is a credibility asset, not a weakness.")

# ---------------------------------------------------------------------
h1("3. Honest Assessment (Brutally Honest Review)")

h2("3.1 What's genuinely strong")
bullet("The core idea is on-target for supply chain analytics / decision science roles, not a generic project wearing a supply-chain costume.")
bullet("The fairness / no-lookahead design is the single best asset in the whole project — a subtle, correct, non-obvious principle most people building a similar comparison would get wrong.")
bullet("Risk-as-a-hard-constraint (rather than a cost penalty folded into the objective) is a strong, well-articulated, specific modeling choice.")
bullet("Splitting into two independent deliverables is the right structural call.")

h2("3.2 Issues identified and how they were resolved")
bullet(
    "The original formulation was labeled a MILP but contained no integer or binary variables "
    "(pure continuous LP). ",
    bold_lead="LP vs. MILP mislabeling — ",
)
p("Resolved: a real minimum-order-quantity (MOQ) constraint was added, which genuinely requires a binary variable (see Section 4.2). This makes \"MILP\" true rather than aspirational.")
bullet(
    "The framing document described \"two bugs caught during validation\" in the past tense, "
    "before any implementation or validation had actually happened. ",
    bold_lead="Bug-story timing — ",
)
p("Resolved: agreed that any bug-catching narrative should be written after real validation happens, describing what actually broke — not pre-scripted before the fact.")
bullet(
    "Model correctness was well-documented, but there was no plan to produce a dollar- or "
    "percentage-quantified takeaway. ",
    bold_lead="Missing quantified takeaway — ",
)
p("Resolved: case study results will be shown via charts quantifying total cost improvement, ordering pattern differences, disruption-response strategy, and inventory behavior (see Section 7).")
bullet(
    "All parameters (costs, tariffs, tiers, capacities) are invented, not calibrated against "
    "real data. ",
    bold_lead="Synthetic data risk — ",
)
p("Resolved: an explicit synthetic-data disclaimer will be placed near the top of the README so it is clear to any reader that the numbers are illustrative, not real-world calibrated.")
bullet(
    "A wall of LaTeX-formatted constraints does not communicate anything at a glance to a "
    "recruiter skimming for under a minute. ",
    bold_lead="Demo-ability gap — ",
)
p("Resolved: results visualizations (tornado chart, cost comparisons, time series) will be front-and-center in the solver README, and a post-completion results document plus a short video playthrough will be produced.")

# ---------------------------------------------------------------------
h1("4. The Optimization Model")

h2("4.1 Core structure (recap)")
bullet("Decision variable: how much each facility orders from each supplier, each period. This is the only thing the solver actually chooses — inventory, backlog, and holding cost are consequences of this choice, not separate free variables.")
bullet("Objective: minimize procurement cost (quantity x unit cost + tariff) plus projected holding cost and projected backorder cost, summed across the full lead-time window — not just the next period.")
bullet("Constraint 1 — Minimum inventory floor, checked at every intermediate checkpoint across the lead-time window (not just at the end of the window). This closes a loophole where the solver could stack orders on the cheapest, slowest supplier and run dangerously low in the middle periods.")
bullet("Constraint 2 — Maximum inventory ceiling, counting only physically-arrived stock (on-hand plus this period's arrivals). Units still in transit do not count against the storage cap.")
bullet("Constraint 3 — Dedicated (non-pooled) capacity per facility-supplier pair, adjusted downward unequally by facility priority rank during disruptions.")
bullet("Constraint 4 — Tiered supplier diversification cap: no single supplier can supply more than a set percentage of a facility's total order in a period, based on that supplier's reliability tier. This is what keeps risk management a visible, inspectable constraint rather than a hidden cost penalty.")
bullet("Constraint 5 — Non-negativity.")
bullet("A separate realized-cost recursion runs after each period's orders are placed, using the actual ground-truth demand once it's revealed. This is kept structurally separate from the decision-time math (which only ever sees the forecast) — this separation is exactly what preserves the fairness / no-lookahead guarantee.")

h2("4.2 Making it a genuine MILP: Minimum Order Quantity (MOQ)")
p(
    "A plain lower-bound constraint (order-quantity greater than or equal to MOQ, unconditionally) "
    "does NOT correctly model a minimum order quantity — it would force every facility to buy at "
    "least the MOQ from every supplier, every single period, whether needed or not, which is unrealistic "
    "and could make scenarios infeasible."
)
p(
    "The real business rule is disjunctive: order nothing, OR order at least the MOQ. The feasible "
    "set for the order-quantity variable becomes {0} union [MOQ, capacity] — a disconnected set with "
    "a gap between zero and the MOQ. A continuous variable governed only by linear inequalities always "
    "has a convex (connected, gap-free) feasible region — this is a basic property of linear "
    "programming. A disconnected feasible set cannot be represented by linear constraints on a "
    "continuous variable alone. A binary variable is required to \"switch\" between the two disconnected "
    "pieces:"
)
p("      MOQ x (engage-binary)  <=  Order Quantity  <=  Capacity x (engage-binary)", italic=True)
p(
    "When the binary is 0, both bounds collapse to force the order quantity to 0. When the binary is 1, "
    "the order quantity is free between the MOQ and capacity. This is a standard OR construct known as "
    "a semi-continuous variable, and it is what genuinely makes this model a MILP rather than an LP — "
    "not a cosmetic addition."
)

h2("4.3 Optional complementary extensions (same binary infrastructure)")
bullet("Fixed ordering / setup cost: add a flat cost incurred any period an order is placed at all, independent of quantity — gives the solver a genuine incentive to batch orders rather than order small amounts every period.")
bullet("Discrete batch / container sizing: order quantity must be an integer multiple of a fixed batch size (pallets/containers) — introduces a genuine general-integer variable, not just binary, modeling realistic packaging/shipping constraints.")

# ---------------------------------------------------------------------
h1("5. Sensitivity Analysis Plan")

h2("5.1 Why classical dual values / shadow prices don't work here")
p(
    "Classical LP duality (shadow prices) is proven on convex feasible regions and relies on the "
    "simplex method's basis structure. Once the MOQ binary variables are added, the true feasible "
    "region is no longer convex, so there is no well-defined \"shadow price\" in the classical sense. "
    "A solver can still return a .pi value on a MILP (typically from the final LP relaxation at the "
    "optimal integer solution), but it does not reliably predict what happens if a parameter is "
    "actually perturbed and the model re-solved — so it should not be trusted or presented as a true "
    "dual value."
)

h2("5.2 The chosen approach: re-optimization / finite-difference sensitivity")
p(
    "Instead of reading dual values off the solver's internals, perturb one parameter at a time and "
    "re-solve the full MILP, measuring the resulting change in total cost. This works identically "
    "whether the model is a pure LP or a MILP, since it never depends on solver-internal duality."
)

h2("5.3 Two flavors of sensitivity")
bullet(
    "Nudge a continuous parameter by a small amount and re-solve, to answer \"what is one unit of "
    "this worth?\" — e.g. supplier capacity, the inventory floor, the warehouse ceiling, the tariff "
    "rate, holding/backorder cost.",
    bold_lead="Marginal sensitivity — ",
)
bullet(
    "Toggle an entire constraint on or off and re-solve, to answer \"what does this whole policy cost "
    "us?\" — e.g. the cost of the tiered diversification requirement, the cost of the MOQ constraint, "
    "the cost of enforcing per-checkpoint inventory feasibility versus only an end-of-window target. "
    "This is the more narratively powerful flavor, since it directly quantifies the dollar value of "
    "design decisions already made in the model (e.g. \"the diversification policy costs $X per period, "
    "in exchange for reduced concentration risk\").",
    bold_lead="Structural / policy sensitivity — ",
)
p(
    "Both flavors should be run across the many-runs automated harness (not a single scenario) and "
    "averaged, to get a robust marginal/policy value rather than a one-off anecdote from a single "
    "random seed."
)

h2("5.4 Output and where it's showcased")
p(
    "Output: a tornado chart (parameters/policies ranked by magnitude of cost impact) plus a small "
    "results table (parameter, perturbation size, cost delta, one-line interpretation)."
)
p(
    "This is a solver-side (Phase 1) deliverable, not something tied to any specific human playthrough. "
    "It is showcased in an \"About the Oracle\" section of the solver's README/results document, framed "
    "as: beyond producing order quantities, the Oracle also surfaces which levers matter most for "
    "future decision-making. This reinforces the \"prescriptive analytics, not just optimization\" "
    "positioning from Section 2.1. It may optionally be echoed as static background content on a "
    "non-interactive \"About\" page in the Phase 2 web app, but must never appear during a "
    "participant's live decision-making flow, to avoid contaminating the data being collected."
)

# ---------------------------------------------------------------------
h1("6. Forecasting Design")

h2("6.1 Menu and rules")
bullet("The forecasting method is chosen once, at the start of a playthrough, and is locked for the entire run — it cannot be changed mid-simulation.")
bullet("A fixed, pre-decided menu of forecasting methods is offered to every participant. Candidate menu: naive (last period), simple moving average at multiple window sizes, weighted moving average, simple exponential smoothing, and optionally trend-adjusted (Holt's) exponential smoothing or linear regression trend forecasting if the demand data has a meaningful trend component.")
bullet("Note: a 1-period moving average is mathematically identical to the naive method — worth merging or keeping deliberately for pedagogical clarity, but not treating as two distinct techniques.")

h2("6.2 How the Oracle chooses, and why it's still fair")
p(
    "The Oracle backtests every method in the same fixed menu against the available historical demand "
    "data and selects whichever produced the lowest historical error, then commits to that choice for "
    "the full run — mirroring the same one-time lock-in rule given to human participants. Because the "
    "Oracle chooses from the identical menu available to participants (no access to a more sophisticated "
    "technique), the identical-information fairness principle from Section 2.2 is preserved: any Oracle "
    "advantage is attributable to disciplined method selection and better decisions, not to a hidden "
    "forecasting capability humans don't have access to."
)

h2("6.3 Gap decomposition: separating forecast quality from decision quality")
p(
    "A three-way comparison isolates how much of a participant's underperformance came from a worse "
    "forecast versus worse ordering decisions:"
)
make_table(
    ["Scenario", "Forecast used", "Decisions used"],
    [
        ["Player (actual)", "Player's chosen forecast", "Player's actual decisions"],
        ["Bridge (hypothetical)", "Player's chosen forecast", "Oracle's optimal decisions"],
        ["Oracle (actual)", "Oracle's chosen forecast", "Oracle's optimal decisions"],
    ],
)
p(
    "Decision-quality gap = Player's realized cost minus Bridge's realized cost (same forecast in both, "
    "isolates decision-making skill)."
)
p(
    "Forecast-quality gap = Bridge's realized cost minus Oracle's realized cost (optimal decisions in "
    "both, isolates forecast accuracy)."
)
p(
    "These two components sum exactly to the total gap between Player and Oracle. This mirrors the "
    "same logic as a price-variance vs. quantity-variance decomposition in standard cost accounting."
)

h2("6.4 Computational structure (cheap, given the locked scenario)")
bullet("The Oracle's own baseline run (its chosen forecast plus optimal decisions) is computed once, total — reused as the benchmark for every participant, since the scenario is identical for everyone.")
bullet("The \"bridge\" run only needs to be computed once per unique forecasting method actually used in the menu (not once per participant), because a given method run against the same shared historical data always produces the identical forecast series. If multiple participants pick the same method, the cached bridge result is reused.")
bullet("Each participant's own actual playthrough is not a solver run at all — it is simply their recorded real decisions.")
p(
    "Given the small problem size, total solver work across the entire case study amounts to roughly "
    "a hundred or so small MILP solves at most — a matter of seconds to a couple of minutes of "
    "computation, not a performance concern."
)

# ---------------------------------------------------------------------
h1("7. Human Case Study Experimental Design")

h2("7.1 The comparability requirement")
p(
    "For results across participants to be interpretable, every participant must face an identical "
    "scenario (demand path, disruption timing, and disruption type). If scenarios differ between "
    "participants, there is no way to tell whether a worse outcome reflects worse decisions or simply "
    "a harder random draw — a classic confound, especially damaging at a small sample size (5-10 "
    "participants) where there isn't enough data to average the variation away."
)

h2("7.2 The chosen design: randomize once, then freeze (\"Interpretation B\")")
p(
    "The number of disruptions is fixed for everyone, but the specific timing and type of each "
    "disruption is generated using a random process exactly once — not hand-scripted by the builder — "
    "and then that single generated scenario is frozen and used identically for every human participant. "
    "This preserves full cross-participant comparability (same design benefits as a fully fixed/scripted "
    "scenario) while avoiding a manually cherry-picked feel, since a random process chose the specifics."
)
p(
    "This is consistent with the standard methodology used in behavioral operations research studies "
    "(e.g. the MIT Beer Distribution Game), where every participant faces an identical, controlled "
    "scenario specifically so that between-participant variation reflects decision quality, not "
    "scenario difficulty."
)

h2("7.3 Getting the \"disruption sensitivity\" finding from a single locked scenario")
p(
    "The one locked scenario should script a few different disruption types at different points in the "
    "timeline (e.g. a supplier capacity cut early, a tariff spike mid-way, a recovery/stable stretch "
    "at the end). Comparing each participant's period-by-period decision-quality gap during disruption "
    "segments versus stable segments, aggregated across all participants, produces the \"does the gap "
    "widen under disruption\" finding without needing multiple scenario variants."
)

h2("7.4 Contrast with the automated many-runs harness")
p(
    "The small-N human case study uses one frozen scenario for everyone. The separate automated "
    "many-runs robustness harness (a rule-based heuristic \"player\" run across roughly 50 random "
    "seeds, compared against the Oracle) is a different design with a different goal, and should "
    "randomize the scenario across runs — since building a distribution across many random conditions "
    "is the point there, not comparing named individuals against each other. The fairness rule still "
    "applies within any single run (the heuristic and the Oracle must see the identical scenario for "
    "that specific run); only across runs does randomization make sense."
)
p("In short: fixed across participants for the human study; randomized across runs for the automated study — fair within any single comparison, either way.", bold=True)

# ---------------------------------------------------------------------
h1("8. Supply Chain Role Coverage")

p(
    "Given the goal of demonstrating breadth across supply chain functions (rather than targeting one "
    "specific role), here is how the current design maps onto standard supply chain function buckets:"
)

make_table(
    ["Function", "Coverage", "Notes"],
    [
        ["Procurement / Sourcing", "Strong", "Tariffs, supplier cost, capacity, MOQ"],
        ["Risk Management", "Strong", "Tiered diversification constraint, disruption modeling"],
        ["Inventory / Replenishment Planning", "Strong", "Floor/ceiling, holding/backorder cost, rolling horizon"],
        ["Operations Research / Analytics", "Strong (core of project)", "The optimization model itself"],
        ["Demand Planning / Forecasting", "Strong", "Fixed forecasting menu, chosen by both player and Oracle; gap decomposition quantifies forecasting's contribution"],
        ["Logistics / Transportation", "Adequate", "Cost/lead-time/reliability tradeoffs already vary naturally across the existing supplier set (domestic vs. international)"],
        ["S&OP", "Weak, cheap to add", "Recommend framing the results document with S&OP-style KPI language (fill rate, inventory turns, total landed cost) alongside the OR results"],
        ["Network Design (strict sense)", "Deliberately excluded", "Facility locations are fixed by design; worth one sentence in the write-up naming this as an explicit scope boundary rather than an oversight"],
    ],
    col_widths=[1.8, 1.3, 3.4],
)

# ---------------------------------------------------------------------
h1("9. Toolchain (Settled)")

h2("9.1 Solver / Phase 1")
bullet("Language/library: Python + PuLP. Chosen over Julia/JuMP for AI-assisted-implementation reliability and to keep one consistent ecosystem across the whole portfolio.")
bullet("Data schema: plain Python dataclasses or a JSON schema, defining Supplier, Facility, Period State, and Disruption Event, so the Oracle's input/output contract is precise and identical for both the player and Oracle sides.")
bullet("Validation: hand-worked test cases asserted against the solver's output — no extra tooling required.")

h2("9.2 Simulation / Phase 2")
bullet("Backend: Python, reusing the same schema, calling the solver as a function or API.")
bullet("Frontend: a simple web page (HTML/JS or React + FastAPI backend) — shareable via link, works with no install, doubles as a demo-video asset.")
bullet("Design prototyping: mock up screens conversationally before committing to the real build.")

# ---------------------------------------------------------------------
h1("10. Simulation Visual/UX Design Principles")

p("Reframed explicitly as a simulation instrument for data collection, not a game:", italic=True)
bullet("Neutral, analytical aesthetic — no score animations, no celebratory feedback, no gamified progress bars, to avoid nudging participant decisions.")
bullet("Consistent layout every period, so variance in the data comes from decisions, not from relearning a shifting interface.")
bullet("No visual hinting toward the Oracle's answer during the decision phase — the Oracle comparison only appears after a decision is committed.")
bullet("Clear, minimal-chrome data displays (plain metric cards/tables), no stylized/decorative charts during play.")
bullet("A short, neutral instructions screen up front, since participants run this without the builder present.")
bullet("A results/comparison screen only at the very end, after all periods are complete.")
p(
    "Data logging requirement: every decision must be logged with at minimum participant_id, period, "
    "facility, supplier, order_quantity, and timestamp, to support cross-participant aggregation."
)

# ---------------------------------------------------------------------
h1("11. Deliverables & Documentation Plan")

bullet("Solver README with the sensitivity tornado chart and results table front-and-center, plus an explicit synthetic-data disclaimer near the top.")
bullet("Case study results, shown through charts: total cost comparison (Oracle vs. heuristic vs. human), ordering pattern overlays, disruption-period cost-gap over time, inventory/backlog time series, and a cost-vs-service-level efficient frontier.")
bullet("A post-completion results document, written only after real validation/playthroughs happen (not pre-scripted), including any bugs actually found.")
bullet("A short video playthrough of the simulation as a demo asset.")

# ---------------------------------------------------------------------
h1("12. Open Items / Still To Decide")

numbered("Repo structure — separate GitHub repos for Solver vs. Simulation, or one monorepo with a hard folder boundary and no cross-imports except through the shared schema.")
numbered("The companion document \"Supply_Chain_Simulation_FULL_BRIEF.md\" (referenced but not yet shared) likely contains the actual numeric parameters and hand-worked test cases needed to implement and validate the solver.")
numbered("Whether to add the fixed-ordering-cost and/or discrete-batch-sizing MILP extensions (Section 4.3) on top of the core MOQ constraint, or keep MOQ as the only integer feature.")
numbered("Whether the demand data has a genuine trend component, which determines whether Holt's trend-adjusted exponential smoothing earns a spot on the forecasting menu (Section 6.1) or is unnecessary complexity.")
numbered("Final confirmation of the rolling-horizon assumption: the Oracle re-solves every period using that period's freshly revealed forecast/state, rather than one single upfront full-horizon solve.")

doc.add_paragraph()
closing = doc.add_paragraph()
closing_run = closing.add_run(
    "This document summarizes the design discussion to date. No implementation has begun; "
    "this is a planning artifact only."
)
closing_run.italic = True
closing_run.bold = True

doc.save("Supply_Chain_Project_Design_Brief.docx")
print("Saved Supply_Chain_Project_Design_Brief.docx")
